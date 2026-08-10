# This file is part of the ECM extension for Indico.
# Copyright (C) 2026 - present
#
# Licensed under the MIT License; see the LICENSE file for details.

"""Producing the documents the provider actually sends.

Two of them: the invitation letter, written into the provider's own Word
template, and the ECM certificate, rendered as a PDF with a verification QR
code.

The letter keeps the existing `.docx`: the layout, the letterhead and the
signature are the provider's, and nothing here rewrites them. Placeholders are
replaced in place, run by run, so formatting survives.
"""

import io
import re

import qrcode

from indico.core.logger import Logger


logger = Logger.get('plugin.ecm.documents')

#: Placeholders as they appear in the Word template: {nomeOspedale}, {{NOME}}, «campo»
PLACEHOLDER_RE = re.compile(r'\{\{?\s*([A-Za-z_][\w°\- ]*)\s*\}?\}|«\s*([A-Za-z_][\w°\- ]*)\s*»')


def _replace_in_paragraph(paragraph, context):
    """Replace placeholders in a paragraph without losing its formatting.

    Word splits a paragraph into runs at arbitrary points, so a placeholder can
    be cut in half. The text is rebuilt once, then written back into the first
    run and the others are emptied — the first run's formatting wins, which is
    what the template intends.
    """
    text = paragraph.text
    if '{' not in text and '«' not in text:
        return False

    def substitute(match):
        name = match.group(1) or match.group(2)
        value = context.get(name)
        if value is None:
            value = context.get(name.strip().replace(' ', '_'))
        return str(value) if value is not None else match.group(0)

    new_text = PLACEHOLDER_RE.sub(substitute, text)
    if new_text == text:
        return False
    for index, run in enumerate(paragraph.runs):
        run.text = new_text if index == 0 else ''
    if not paragraph.runs:
        paragraph.add_run(new_text)
    return True


def render_docx(template_path, context):
    """Fill the Word template and return the document as bytes."""
    from docx import Document

    document = Document(str(template_path))
    replaced = 0
    for paragraph in document.paragraphs:
        replaced += _replace_in_paragraph(paragraph, context)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replaced += _replace_in_paragraph(paragraph, context)
    for section in document.sections:
        for part in (section.header, section.footer):
            for paragraph in part.paragraphs:
                replaced += _replace_in_paragraph(paragraph, context)

    buffer = io.BytesIO()
    document.save(buffer)
    logger.debug('rendered %s with %d replacements', template_path, replaced)
    return buffer.getvalue()


def missing_placeholders(template_path, context):
    """Placeholders the template uses and the context does not provide.

    Used before a batch: a letter that goes out with `{nomeOspedale}` printed on
    it is worse than a letter that was never generated.
    """
    from docx import Document

    document = Document(str(template_path))
    texts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        texts.extend(cell.text for row in table.rows for cell in row.cells)
    missing = set()
    for text in texts:
        for match in PLACEHOLDER_RE.finditer(text):
            name = match.group(1) or match.group(2)
            if name not in context:
                missing.add(name)
    return sorted(missing)


def qr_png(data, *, box_size=6, border=2):
    """A QR code as PNG bytes."""
    code = qrcode.QRCode(box_size=box_size, border=border,
                         error_correction=qrcode.constants.ERROR_CORRECT_M)
    code.add_data(data)
    code.make(fit=True)
    buffer = io.BytesIO()
    code.make_image(fill_color='black', back_color='white').save(buffer, format='PNG')
    return buffer.getvalue()


def render_pdf(html, *, base_url=None):
    """Render HTML into PDF bytes with WeasyPrint, as Indico does for receipts."""
    from weasyprint import HTML

    return HTML(string=html, base_url=base_url).write_pdf()
