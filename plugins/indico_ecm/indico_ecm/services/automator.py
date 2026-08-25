# This file is part of the ECM extension for Indico.
# Copyright (C) 2026 - present
#
# Licensed under the MIT License; see the LICENSE file for details.

"""The automator.

In the legacy application this is the feature that reads the email and the
attachments a sponsor sends and sets up the event: it extracts the details,
proposes a folder name and drafts the starting documents.

Two things change in the port, and both are the point of it:

- the prompt is **data**, versioned and hashed, so what an extraction was asked
  to do is recoverable months later;
- the deterministic part is separated from the model. Dates, event codes and
  folder names are found with rules; the model is only needed for prose. When
  the rules are enough, no model is called at all.

Nothing here calls a language model: this module prepares the request and
validates the answer. The call belongs to the agent runtime, behind its
permission table.

Pure, no Indico imports.
"""

import hashlib
import re
from dataclasses import dataclass, field
from datetime import date

from indico_ecm.services.legacy_import import parse_date
from indico_ecm.services.naming import generate_folder_name
from indico_ecm.services import programme
from indico_ecm.services.specialty import identify_event_format, identify_specialty


#: The instruction ported from `state.js`, kept as the provider wrote it.
#: Change it by adding a new version, never by editing this string in place.
AUTOMATOR_PROMPT_V1 = '''## Role and Goal
You are an expert event management assistant for a medical event organizer. Your goal is to analyze
provided email text and documents to automatically set up a new event. You must extract key
information, generate the starting documents, and prepare the folder structure.

## Analysis and Extraction
1. Analyze: read the user-provided email text and any attached documents.
2. Extract: eventName, eventDate (YYYY-MM-DD), eventLocation, eventSponsor, eventCode, speakers.
3. Folder name: follow the organizer's convention, which is provided to you already computed.
   Do not invent a different one.

## Content Generation
Produce professional, concise content for: info_evento.txt, briefing.txt, agenda.txt,
report_template.txt, email_draft.html.

## Rules
- Do not invent dates, codes or names that are not in the source material. Leave a field empty
  instead, and say what was missing.
- Do not state the number of ECM credits: credits are decided by the accreditation dossier.
- Report every value you extracted together with the sentence you took it from.

## Output Format
Return a single JSON object adhering to the provided schema, with no text outside it.'''

AUTOMATOR_PROMPT_VERSION = 'automator-v1'

#: The response contract, ported from `automatorResponseSchema`
AUTOMATOR_RESPONSE_SCHEMA = {
    'type': 'object',
    'required': ['extractedData', 'fileContents'],
    'properties': {
        'folderName': {'type': 'string'},
        'extractedData': {
            'type': 'object',
            'properties': {
                'eventName': {'type': 'string'},
                'eventDate': {'type': 'string'},
                'eventLocation': {'type': 'string'},
                'eventSponsor': {'type': 'string'},
                'eventCode': {'type': 'string'},
                'speakers': {'type': 'array', 'items': {'type': 'string'}},
            },
        },
        'fileContents': {
            'type': 'object',
            'properties': {
                'info_evento.txt': {'type': 'string'},
                'briefing.txt': {'type': 'string'},
                'agenda.txt': {'type': 'string'},
                'report_template.txt': {'type': 'string'},
                'email_draft.html': {'type': 'string'},
            },
        },
        'evidence': {
            'type': 'object',
            'description': 'For each extracted field, the sentence it came from.',
        },
        'imagePrompt': {'type': 'string'},
    },
}


def prompt_fingerprint(prompt=AUTOMATOR_PROMPT_V1):
    """Hash of the instruction, recorded on every extraction."""
    return hashlib.sha256(prompt.encode()).hexdigest()[:16]


#: The provider's own convention: 0116_GDBO, 2026-CARD. Deliberately narrow —
#: a looser pattern matches VAT numbers, room names and invoice references.
EVENT_CODE_RE = re.compile(r'\b(\d{3,4}[_-][A-Z]{2,6}|\d{4}-[A-Z]{2,6})\b')
#: Words that announce a code, after which almost anything is one
CODE_KEYWORD_RE = re.compile(r'(?:codice\s+evento|codice|cod\.|rif\.|rif\b)\s*[:\-\u2013]?', re.IGNORECASE)
CODE_AFTER_KEYWORD_RE = re.compile(r'\s*([A-Za-z0-9][A-Za-z0-9_\-/]{2,19})\b')

DATE_RE = re.compile(r'\b(\d{1,2}[/.-]\d{1,2}[/.-]\d{2,4}|\d{4}-\d{2}-\d{2})\b')

#: "Dott. Mario Rossi", "Prof.ssa Anna Verdi", "Prof. Gian Luca De Angelis"
_TITLE = r'(?:Dott\.ssa|Dott\.|Prof\.ssa|Prof\.|Dr\.ssa|Dr\.)'
_NAME = r"[A-ZÀ-Ú][\w'\u2019À-ú]+"
#: Italian and other European surname particles, which are lowercase and would
#: otherwise cut a surname in half ("De Angelis" -> "De")
_PARTICLE = r"(?:de(?:gli|lla|llo|lle|l|i|n|r)?|di|da|dal|del|della|van|von|ten|ter|la|lo|d')"
_TOKEN = rf'(?:{_NAME}|{_PARTICLE})\b'
SPEAKER_RE = re.compile(rf'\b{_TITLE}\s*({_TOKEN}(?:\s+{_TOKEN}){{0,3}})')
#: Role words that follow a name and are not part of it
ROLE_WORDS = frozenset({'presidente', 'direttore', 'responsabile', 'coordinatore', 'moderatore',
                        'relatore', 'relatrice', 'segretario', 'segretaria', 'docente', 'tutor',
                        # A programme puts these labels on the line after a name;
                        # without them a speaker came out as "Mario Rossi Faculty".
                        'faculty', 'scientifico', 'scientifica', 'discenti', 'partecipanti'})


def find_event_code(text):
    """Find the event code, preferring one that is announced as such.

    A code introduced by "codice" or "rif." is taken as given; otherwise only
    the provider's own format is accepted. Anything looser matches a VAT number
    or a room name, and a wrong code sends the accreditation request to the
    wrong folder.
    """
    for keyword in CODE_KEYWORD_RE.finditer(text or ''):
        match = CODE_AFTER_KEYWORD_RE.match(text, keyword.end())
        if match and any(char.isdigit() for char in match.group(1)):
            return match.group(1)
    match = EVENT_CODE_RE.search(text or '')
    return match.group(1) if match else ''


def find_speakers(text):
    """Names introduced by an academic or medical title.

    Trailing role words are dropped: "Dott. Mario Rossi Presidente" is a person
    called Mario Rossi, not one called Mario Rossi Presidente.
    """
    speakers = []
    for raw in SPEAKER_RE.findall(text or ''):
        words = raw.split()
        while words and words[-1].casefold() in ROLE_WORDS:
            words.pop()
        name = ' '.join(words).strip()
        if name and name not in speakers:
            speakers.append(name)
    return speakers


@dataclass
class Extraction:
    """What the deterministic pass could establish on its own."""

    event_name: str = ''
    event_date: date | None = None
    end_date: date | None = None
    location: str = ''
    sponsor: str = ''
    event_code: str = ''
    speakers: list = field(default_factory=list)
    specialty: str = ''
    activity_format: str = ''
    folder_name: str = ''
    venue: str = ''
    city: str = ''
    province: str = ''
    credits: str = ''
    participants: str = ''
    start_time: str = ''
    end_time: str = ''
    #: The speakers with the role and title their line carried
    people: list = field(default_factory=list)
    #: field -> the sentence it was taken from
    evidence: dict = field(default_factory=dict)
    #: What the rules could not determine and a model (or a person) must supply
    unresolved: list = field(default_factory=list)

    @property
    def needs_model(self):
        return bool(self.unresolved)


def _sentence_containing(text, needle):
    for sentence in re.split(r'(?<=[.!?\n])\s+', text or ''):
        if needle and needle in sentence:
            return sentence.strip()[:300]
    return ''


def extract(text, *, known_sponsor='', known_location=''):
    """Find what can be found with rules alone.

    Deliberately conservative: it reports what it saw and where, and lists what
    it could not resolve, rather than producing a confident guess.

    The document itself is read by `services/programme.py`, which knows the
    shape of a Progetto Formativo — the title under the header, the labelled
    lines, the timetable with the speakers under each session. What stays here
    is what is about the event rather than the document: the provider's own
    event code, the specialty, the format.
    """
    text = text or ''
    result = Extraction(sponsor=known_sponsor, location=known_location)
    document = programme.read(text)
    result.evidence.update(document.evidence)

    code = find_event_code(text)
    if code:
        result.event_code = code
        result.evidence['event_code'] = _sentence_containing(text, code)
    else:
        result.unresolved.append('event_code')

    if document.start_date:
        result.event_date = document.start_date
        result.end_date = document.end_date
    else:
        result.unresolved.append('event_date')

    # A title in front of a name is the strongest signal; the programme reader
    # adds the ones written without one, under their session.
    speakers = list(find_speakers(text))
    for person in document.people:
        if person.name not in speakers:
            speakers.append(person.name)
    result.speakers = speakers
    result.people = document.people
    if not speakers:
        result.unresolved.append('speakers')

    if document.event_name:
        result.event_name = document.event_name
    else:
        result.unresolved.append('event_name')

    if not result.location:
        result.location = document.city or document.venue
    result.venue = document.venue
    result.city = document.city
    result.province = document.province
    result.credits = document.credits
    result.participants = document.participants
    result.start_time = document.start_time
    result.end_time = document.end_time

    match = identify_specialty(text)
    result.specialty = match.specialty
    result.activity_format = identify_event_format(text)

    if not result.location:
        result.unresolved.append('location')
    if not result.sponsor:
        result.unresolved.append('sponsor')
    if not result.activity_format:
        result.unresolved.append('activity_format')
    return result


def folder_name_for(extraction: Extraction, *, event_name='', city=''):
    """The folder name, computed by the platform rather than by the model."""
    return generate_folder_name(
        start_date=extraction.event_date,
        end_date=extraction.end_date,
        event_name=event_name or extraction.event_name,
        event_type=extraction.activity_format,
        city=city or extraction.location,
        sponsor=extraction.sponsor,
        event_code=extraction.event_code,
    )


def build_request(text, *, known_sponsor='', known_location=''):
    """Everything the agent runtime needs to ask a model for the rest.

    The folder name is passed in already computed: the convention belongs to the
    provider, not to whatever the model would come up with.
    """
    extraction = extract(text, known_sponsor=known_sponsor, known_location=known_location)
    return {
        'prompt': AUTOMATOR_PROMPT_V1,
        'prompt_version': AUTOMATOR_PROMPT_VERSION,
        'prompt_sha': prompt_fingerprint(),
        'response_schema': AUTOMATOR_RESPONSE_SCHEMA,
        'deterministic': {
            'event_code': extraction.event_code,
            'event_date': extraction.event_date.isoformat() if extraction.event_date else '',
            'speakers': list(extraction.speakers),
            'specialty': extraction.specialty,
            'activity_format': extraction.activity_format,
            'folder_name': folder_name_for(extraction),
        },
        'unresolved': list(extraction.unresolved),
        'source_length': len(text or ''),
    }


class AutomatorError(Exception):
    pass


def validate_response(payload, *, deterministic=None):
    """Check a model answer before anything is created from it.

    The deterministic values win: if the model returns a different event code or
    folder name, the platform's own value is kept and the disagreement is
    reported. A model does not get to rename a folder that already has a
    convention.
    """
    if not isinstance(payload, dict):
        raise AutomatorError('risposta non valida: atteso un oggetto JSON')
    for key in ('extractedData', 'fileContents'):
        if key not in payload:
            raise AutomatorError(f'risposta incompleta: manca {key}')

    extracted = dict(payload.get('extractedData') or {})
    files = dict(payload.get('fileContents') or {})
    conflicts = []
    deterministic = deterministic or {}

    for model_key, own_key in (('eventCode', 'event_code'), ('folderName', 'folder_name')):
        own_value = deterministic.get(own_key)
        model_value = extracted.get(model_key) if model_key == 'eventCode' else payload.get('folderName')
        if own_value and model_value and str(model_value).strip() != str(own_value).strip():
            conflicts.append({'field': own_key, 'platform': own_value, 'model': model_value})

    if deterministic.get('event_code'):
        extracted['eventCode'] = deterministic['event_code']
    folder_name = deterministic.get('folder_name') or payload.get('folderName', '')

    return {
        'extracted': extracted,
        'files': {name: content for name, content in files.items() if content},
        'folder_name': folder_name,
        'conflicts': conflicts,
        'image_prompt': payload.get('imagePrompt', ''),
        'evidence': payload.get('evidence', {}),
    }


#: What the starting documents of an event folder contain.
#: Deliberately produced from the extracted data, not from a model: an empty
#: section a person fills in is better than a plausible paragraph nobody wrote.
FOLDER_TEMPLATES = {
    'info_evento.txt': '''EVENTO
======
Nome:        {event_name}
Codice:      {event_code}
Data:        {dates}
Luogo:       {place}
Sponsor:     {sponsor}
Modalità:    {activity_format}
Specialità:  {specialty}

RELATORI
--------
{speakers}

NOTE
----
{unresolved_note}
''',
    'briefing.txt': '''BRIEFING INTERNO — {event_name}
{underline}

Data:     {dates}
Luogo:    {place}
Sponsor:  {sponsor}
Codice:   {event_code}

RUOLI
-----
Responsabile scientifico:
Project manager:
Segreteria:
Grafica:
Hostess:

SCADENZE
--------
Accreditamento:
Contratti sponsor:
Grafica:
Lettere di incarico:
Slide kit:

LOGISTICA
---------
Sede:
Catering:
Hotel:
Materiali:
''',
    'agenda.txt': '''AGENDA — {event_name}
{underline}
{dates} — {place}

  ora    argomento                              relatore
  -----  -------------------------------------  ------------------------
{agenda_rows}
''',
    'report_template.txt': '''REPORT POST EVENTO — {event_name}
{underline}

PARTECIPAZIONE
--------------
Iscritti:
Presenti:
Aventi diritto ai crediti:
Attestati emessi:

VALUTAZIONE
-----------
Questionari raccolti:
Gradimento medio:
Esito valutazione apprendimento:

OSSERVAZIONI
------------

AZIONI PER LA PROSSIMA EDIZIONE
-------------------------------
''',
    'email_draft.html': '''<p>Gentile Dottoressa, Gentile Dottore,</p>
<p>siamo lieti di invitarLa all'evento <strong>{event_name}</strong>,
che si terrà {dates}{place_clause}.</p>
<p>{specialty_clause}</p>
<p>Per iscriversi: <a href="#">link di iscrizione</a></p>
<p>Cordiali saluti,<br>La segreteria organizzativa</p>
''',
}


def _folder_context(extraction, *, event_name='', sponsor='', place=''):
    name = event_name or extraction.event_name or 'Evento senza titolo'
    dates = ''
    if extraction.event_date:
        dates = extraction.event_date.strftime('%d/%m/%Y')
        if extraction.end_date and extraction.end_date != extraction.event_date:
            dates += f" - {extraction.end_date.strftime('%d/%m/%Y')}"
    speakers = '\n'.join(f'- {speaker}' for speaker in extraction.speakers) or '- (da confermare)'
    agenda_rows = '\n'.join(f'  {"":5}  {"":37}  {speaker}' for speaker in extraction.speakers)
    unresolved = ', '.join(extraction.unresolved)
    return {
        'event_name': name,
        'underline': '=' * min(len(name) + 20, 70),
        'event_code': extraction.event_code or '(da assegnare)',
        'dates': dates or '(da confermare)',
        'place': place or extraction.location or '(da confermare)',
        'place_clause': f' presso {place or extraction.location}' if (place or extraction.location) else '',
        'sponsor': sponsor or extraction.sponsor or '(da confermare)',
        'activity_format': extraction.activity_format,
        'specialty': extraction.specialty,
        'specialty_clause': f'Area tematica rilevata: {extraction.specialty}.' if extraction.specialty else '',
        'speakers': speakers,
        'agenda_rows': agenda_rows or '  (da compilare)',
        'unresolved_note': (f'Dati non ricavati dal materiale: {unresolved}.'
                            if unresolved else 'Tutti i dati sono stati ricavati dal materiale.'),
    }


def build_folder_files(extraction, *, event_name='', sponsor='', place=''):
    """The starting documents of an event folder, as `{filename: text}`."""
    context = _folder_context(extraction, event_name=event_name, sponsor=sponsor, place=place)
    return {name: template.format(**context) for name, template in FOLDER_TEMPLATES.items()}


def build_folder_archive(extraction, *, event_name='', sponsor='', place='', city='', extra_files=None):
    """Build the event folder as a zip, named by the provider's convention.

    Returns `(folder_name, zip_bytes)`. The archive contains one directory named
    after the folder, so unzipping it on the shared drive puts everything in the
    right place in one step.
    """
    import io
    import zipfile

    folder = folder_name_for(extraction, event_name=event_name, city=city or place) or 'EVENTO'
    files = build_folder_files(extraction, event_name=event_name, sponsor=sponsor, place=place)
    files |= (extra_files or {})
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, 'w', zipfile.ZIP_DEFLATED) as archive:
        for name, content in files.items():
            data = content.encode('utf-8') if isinstance(content, str) else content
            archive.writestr(f'{folder}/{name}', data)
    return folder, buffer.getvalue()


def read_document(content, filename=''):
    """Extract text from an uploaded file.

    Handles what actually arrives from a sponsor: plain text, Word and PDF.
    Anything else is reported rather than guessed at.
    """
    lowered = (filename or '').lower()
    if lowered.endswith(('.txt', '.md', '.csv')):
        return content.decode('utf-8', 'replace')
    if lowered.endswith(('.docx', '.docm')):
        import io

        from docx import Document
        document = Document(io.BytesIO(content))
        parts = [paragraph.text for paragraph in document.paragraphs]
        parts += [cell.text for table in document.tables for row in table.rows for cell in row.cells]
        return '\n'.join(part for part in parts if part.strip())
    if lowered.endswith('.pdf'):
        import io

        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(content))
        return '\n'.join((page.extract_text() or '') for page in reader.pages)
    if lowered.endswith(('.eml', '.msg', '.html', '.htm')):
        text = content.decode('utf-8', 'replace')
        return re.sub(r'<[^>]+>', ' ', text)
    raise ValueError(f'formato non supportato: {filename or "file senza nome"}')
